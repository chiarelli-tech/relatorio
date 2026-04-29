"""
Dashboard RDP - Convergint
Gerador automático - lê arquivos locais e gera HTML com dados embutidos.
Execute diariamente via Task Scheduler ou duplo clique no .bat
"""

import os
import re
import json
import sys
import subprocess
from datetime import datetime, date, timedelta
from pathlib import Path

# ── Configurações de caminho ────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
PESSOAS_XLSX   = BASE_DIR / "Pessoas.xlsx"
RDP_DIR        = BASE_DIR / "RDP"
EMAIL_XLSX     = RDP_DIR  / "info_email.xlsx"
OUTPUT_HTML         = BASE_DIR / "dashboard_rdp.html"
OUTPUT_PENDENCIAS   = BASE_DIR / "dashboard_pendencias.html"
PENDENCIAS_JSON     = BASE_DIR / ".pendencias.json"   # persiste pendências entre execuções

# ── Instala dependências se necessário ─────────────────────────────────────
def instalar_deps():
    pkgs = ["openpyxl", "python-docx"]
    for pkg in pkgs:
        try:
            __import__(pkg.replace("-", "_"))
        except ImportError:
            print(f"Instalando {pkg}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "--quiet"])

instalar_deps()

import openpyxl

# ── Helpers ─────────────────────────────────────────────────────────────────
FILENAME_RE = re.compile(
    r'^([A-Z0-9]+)_(\d{4}-\d{2}-\d{2})_RDP_\(([^)]+)\)_(.+)\.(docx?|pdf)$',
    re.IGNORECASE
)

def parse_filename(fname):
    """Extrai metadados do nome do arquivo RDP."""
    m = FILENAME_RE.match(fname)
    if not m:
        return None
    return {
        "project":      m.group(1).upper(),
        "date":         m.group(2),
        "discipline":   m.group(3).upper(),
        "professional": m.group(4),
        "ext":          m.group(5).lower(),
        "filename":     fname
    }

def extract_docx_text(path: Path) -> str:
    """Extrai texto de um .docx."""
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Extrai também tabelas
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs[:80])  # máx 80 linhas
    except Exception as e:
        return ""

def extract_pdf_text(path: Path) -> str:
    """Extrai texto de um PDF."""
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            lines = []
            for page in pdf.pages[:5]:
                t = page.extract_text()
                if t:
                    lines.extend(t.splitlines())
            return "\n".join(l.strip() for l in lines if l.strip())[:4000]
    except Exception:
        pass
    try:
        import PyPDF2
        with open(str(path), "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages[:5]:
                text += page.extract_text() or ""
            return text[:4000]
    except Exception:
        return ""

def extract_activities_and_pending(text: str, source_info: dict) -> dict:
    """
    Analisa o texto do RDP para extrair atividades realizadas e pendências.
    Retorna {"activities": [...], "pending": [...]}
    """
    activities = []
    pending    = []

    if not text:
        return {"activities": activities, "pending": pending}

    lines = [l.strip() for l in text.splitlines() if l.strip()]

    # Heurísticas para detectar seções de atividades e pendências
    ATIV_MARKERS   = ["atividade", "executado", "realizado", "serviço", "serviços", "tarefa", "trabalho"]
    PEND_MARKERS   = ["pendência", "pendencia", "pendente", "aguardando", "falta", "a fazer", "próximo", "proximo", "solicitação"]
    SKIP_MARKERS   = ["att", "atenciosamente", "convergint", "programador", "atenção"]

    in_activities = False
    in_pending    = False

    for line in lines:
        ll = line.lower()

        # Detecta início de seção
        if any(m in ll for m in ATIV_MARKERS) and len(line) < 60:
            in_activities = True
            in_pending    = False
            continue
        if any(m in ll for m in PEND_MARKERS) and len(line) < 60:
            in_pending    = True
            in_activities = False
            continue
        if any(m in ll for m in SKIP_MARKERS):
            in_activities = False
            in_pending    = False
            continue

        # Captura linhas de conteúdo
        if len(line) > 10:
            if in_activities:
                activities.append(line)
            elif in_pending:
                pending.append(line)
            # Se ainda não entrou em seção, tenta identificar pelo conteúdo
            elif any(m in ll for m in PEND_MARKERS):
                pending.append(line)

    return {
        "activities": activities[:20],
        "pending":    pending[:10]
    }

# ── Leitura de Pessoas ───────────────────────────────────────────────────────
def load_pessoas() -> list:
    if not PESSOAS_XLSX.exists():
        print(f"AVISO: {PESSOAS_XLSX} não encontrado.")
        return []
    wb = openpyxl.load_workbook(PESSOAS_XLSX)
    ws = wb.active
    headers = [str(c.value or "").strip().lower() for c in ws[1]]
    pessoas = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        nome  = row[next((i for i,h in enumerate(headers) if "nome" in h), -1)]
        email = row[next((i for i,h in enumerate(headers) if "e-mail" in h or "email" in h), -1)]
        equipe = row[next((i for i,h in enumerate(headers) if "equipe" in h), -1)]
        nome  = str(nome or "").strip()
        email = str(email or "").strip().lower()
        equipe = str(equipe or "").strip()
        if nome and email and "@" in email:
            pessoas.append({"nome": nome, "email": email, "equipe": equipe})
    print(f"  ✓ {len(pessoas)} colaboradores carregados")
    return pessoas

# ── Leitura de Emails ────────────────────────────────────────────────────────
def load_emails() -> list:
    if not EMAIL_XLSX.exists():
        print(f"AVISO: {EMAIL_XLSX} não encontrado.")
        return []
    wb = openpyxl.load_workbook(EMAIL_XLSX)
    ws = wb.active
    headers = [str(c.value or "").strip().lower() for c in ws[1]]

    emails = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        def get(key):
            idx = next((i for i,h in enumerate(headers) if key in h), -1)
            return str(row[idx] or "").strip() if idx >= 0 and row[idx] is not None else ""

        email_rem = get("emailremetente") or get("email")
        if not email_rem or "@" not in email_rem:
            continue

        data_raw = get("datarecebimento")
        data_str = data_raw[:10] if data_raw else ""

        arquivo = get("quantidadeanexos") or get("nomeanexo") or ""
        parsed  = parse_filename(arquivo) if arquivo else None

        emails.append({
            "nomeRemetente":      get("nomeremetente"),
            "emailRemetente":     email_rem.lower(),
            "dataRecebimento":    data_raw,
            "dataStr":            data_str,
            "assunto":            get("assunto"),
            "arquivo":            arquivo,
            "parsed":             parsed,
            "bodyPreview":        get("bodypreview"),
            "statusProcessamento": get("statusprocessamento"),
            "caminhoPasta":       get("caminhopasta"),
        })
    print(f"  ✓ {len(emails)} e-mails carregados")
    return emails

# ── Escaneia pastas de RDP ───────────────────────────────────────────────────
def scan_rdp_folders(pessoas: list) -> list:
    """
    Varre RDP/<YYYY-MM-DD>/ buscando .docx e .pdf.
    Lê conteúdo e extrai atividades/pendências.
    """
    docs = []
    if not RDP_DIR.exists():
        return docs

    email_by_firstname = {}
    for p in pessoas:
        first = p["nome"].split()[0].lower()
        email_by_firstname.setdefault(first, p["email"])

    for entry in sorted(RDP_DIR.iterdir()):
        if not (entry.is_dir() and re.match(r'\d{4}-\d{2}-\d{2}', entry.name)):
            continue
        folder_date = entry.name
        for fpath in entry.iterdir():
            if fpath.suffix.lower() not in (".docx", ".doc", ".pdf"):
                continue

            meta = parse_filename(fpath.name)
            if not meta:
                # Tenta extrair informações mínimas do nome
                meta = {
                    "project": "DESCONHECIDO",
                    "date": folder_date,
                    "discipline": "—",
                    "professional": fpath.stem,
                    "ext": fpath.suffix.lstrip(".").lower(),
                    "filename": fpath.name
                }

            # Extrai texto do documento
            if fpath.suffix.lower() in (".docx", ".doc"):
                text = extract_docx_text(fpath)
            else:
                text = extract_pdf_text(fpath)

            extracted = extract_activities_and_pending(text, meta)

            # Descobre e-mail do profissional
            prof_first = meta["professional"].split()[0].lower()
            prof_email = email_by_firstname.get(prof_first, "")
            for p in pessoas:
                if meta["professional"].lower() in p["nome"].lower() or p["nome"].lower().startswith(prof_first):
                    prof_email = p["email"]
                    break

            docs.append({
                **meta,
                "folderDate":   folder_date,
                "profEmail":    prof_email,
                "activities":   extracted["activities"],
                "pendingItems": extracted["pending"],
                "hasText":      bool(text),
                "textPreview":  text[:300] if text else ""
            })

    print(f"  ✓ {len(docs)} documentos RDP escaneados")
    return docs

# ── Pendências persistentes ──────────────────────────────────────────────────
def load_pendencias() -> list:
    if PENDENCIAS_JSON.exists():
        try:
            return json.loads(PENDENCIAS_JSON.read_text(encoding="utf-8"))
        except Exception:
            pass
    return []

def save_pendencias(pendencias: list):
    PENDENCIAS_JSON.write_text(
        json.dumps(pendencias, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )

def auto_update_pendencias(pendencias: list, docs: list) -> list:
    """
    Lógica automática:
    - Pendências extraídas dos documentos são adicionadas se novas
    - Pendências que não aparecem mais nos últimos 7 dias de docs são marcadas como Concluídas
    """
    today = date.today().isoformat()
    cutoff = (date.today() - timedelta(days=7)).isoformat()

    # Coleta pendências recentes dos documentos
    recent_texts = set()
    for doc in docs:
        if doc.get("folderDate", "") >= cutoff:
            for pi in doc.get("pendingItems", []):
                recent_texts.add(pi[:60].lower().strip())

    # Verifica pendências existentes - marca concluídas se sumiram dos RDPs
    for p in pendencias:
        if p.get("status") == "Aberta" and p.get("autoDetected"):
            desc_key = p.get("desc", "")[:60].lower().strip()
            if desc_key and desc_key not in recent_texts:
                p["status"] = "Concluída"
                p["doneDate"] = today
                p["autoConcluded"] = True

    # Adiciona novas pendências extraídas dos docs que ainda não existem
    existing_descs = {p.get("desc", "")[:60].lower().strip() for p in pendencias}
    for doc in docs:
        for pi in doc.get("pendingItems", []):
            key = pi[:60].lower().strip()
            if key and key not in existing_descs:
                pendencias.append({
                    "id":           int(datetime.now().timestamp() * 1000) + len(pendencias),
                    "desc":         pi,
                    "project":      doc.get("project", ""),
                    "discipline":   doc.get("discipline", ""),
                    "priority":     "Média",  # padrão
                    "status":       "Aberta",
                    "responsible":  doc.get("profEmail", ""),
                    "forecast":     "",
                    "doneDate":     "",
                    "origin":       f"RDP de {doc['professional']} – {doc['folderDate']}",
                    "createdAt":    today,
                    "autoDetected": True
                })
                existing_descs.add(key)

    return pendencias

# ── Geração do HTML ──────────────────────────────────────────────────────────
def gerar_html(pessoas, emails, docs, pendencias):
    today_str = date.today().isoformat()
    generated_at = datetime.now().strftime("%d/%m/%Y às %H:%M")

    # Serializa dados para JSON embutido no HTML
    data_json = json.dumps({
        "pessoas":    pessoas,
        "emails":     emails,
        "docs":       docs,
        "pendencias": pendencias,
        "generatedAt": generated_at,
        "today":      today_str
    }, ensure_ascii=False, default=str)

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Dashboard RDP — Convergint</title>
<script src="https://cdnjs.cloudflare.com/ajax/libs/xlsx/0.18.5/xlsx.full.min.js"></script>
<script src="https://cdnjs.cloudflare.com/ajax/libs/Chart.js/4.4.0/chart.umd.min.js"></script>
<style>
:root{{
  --primary:#0066CC;--primary-dark:#004A99;--success:#16a34a;--danger:#dc2626;
  --warning:#d97706;--info:#0891b2;--bg:#f1f5f9;--card:#ffffff;
  --border:#e2e8f0;--text:#1e293b;--text-muted:#64748b;--sidebar-w:230px;
}}
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Segoe UI',system-ui,sans-serif;background:var(--bg);color:var(--text);display:flex;min-height:100vh}}
#sidebar{{width:var(--sidebar-w);background:var(--primary-dark);color:#fff;display:flex;flex-direction:column;position:fixed;top:0;left:0;height:100vh;z-index:100;overflow-y:auto}}
.logo{{padding:18px 16px 14px;border-bottom:1px solid rgba(255,255,255,.15)}}
.logo span{{display:block;font-size:18px;font-weight:800;margin-bottom:1px}}
.logo small{{font-size:11px;opacity:.6;text-transform:uppercase;letter-spacing:.5px}}
#sidebar nav a{{display:flex;align-items:center;gap:10px;padding:10px 18px;color:rgba(255,255,255,.72);text-decoration:none;font-size:13.5px;font-weight:500;border-left:3px solid transparent;transition:all .15s}}
#sidebar nav a:hover{{color:#fff;background:rgba(255,255,255,.1)}}
#sidebar nav a.active{{color:#fff;background:rgba(255,255,255,.15);border-left-color:#60a5fa}}
#sidebar nav a .icon{{font-size:16px;width:20px;text-align:center}}
.sidebar-footer{{padding:12px 16px;border-top:1px solid rgba(255,255,255,.12);font-size:11px;color:rgba(255,255,255,.5);line-height:1.6}}
#main{{margin-left:var(--sidebar-w);flex:1;display:flex;flex-direction:column;min-height:100vh}}
#topbar{{background:var(--card);border-bottom:1px solid var(--border);padding:11px 24px;display:flex;align-items:center;justify-content:space-between;position:sticky;top:0;z-index:50}}
#topbar h1{{font-size:17px;font-weight:700}}
.actions{{display:flex;gap:10px;align-items:center}}
.btn{{padding:7px 14px;border-radius:6px;border:none;cursor:pointer;font-size:13px;font-weight:600;transition:all .15s;display:inline-flex;align-items:center;gap:6px}}
.btn-primary{{background:var(--primary);color:#fff}}.btn-primary:hover{{background:var(--primary-dark)}}
.btn-outline{{background:transparent;border:1.5px solid var(--border);color:var(--text)}}.btn-outline:hover{{background:var(--bg)}}
.btn-sm{{padding:4px 10px;font-size:12px}}
.btn-success{{background:var(--success);color:#fff}}
.btn-danger{{background:var(--danger);color:#fff;border-color:var(--danger)}}
#content{{padding:22px 24px;flex:1}}
.section{{display:none}}.section.active{{display:block}}
.section-title{{font-size:20px;font-weight:800;margin-bottom:3px}}
.section-sub{{font-size:13px;color:var(--text-muted);margin-bottom:18px}}
.kpi-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(160px,1fr));gap:14px;margin-bottom:22px}}
.kpi-card{{background:var(--card);border-radius:10px;padding:16px 18px;border:1px solid var(--border);position:relative;overflow:hidden}}
.kpi-card::after{{content:'';position:absolute;top:0;left:0;width:4px;height:100%;background:var(--accent,var(--primary))}}
.kpi-card .label{{font-size:11px;color:var(--text-muted);font-weight:700;text-transform:uppercase;letter-spacing:.5px;margin-bottom:5px}}
.kpi-card .value{{font-size:28px;font-weight:800;line-height:1}}
.kpi-card .sub{{font-size:11px;color:var(--text-muted);margin-top:3px}}
.card{{background:var(--card);border-radius:10px;border:1px solid var(--border);overflow:hidden;margin-bottom:18px}}
.card-header{{padding:13px 18px;border-bottom:1px solid var(--border);display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:8px}}
.card-header h2{{font-size:14px;font-weight:700}}
.card-body{{padding:0;overflow-x:auto}}
table{{width:100%;border-collapse:collapse;font-size:13px}}
th{{background:#f8fafc;padding:9px 14px;text-align:left;font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:.4px;color:var(--text-muted);border-bottom:1px solid var(--border);white-space:nowrap}}
td{{padding:9px 14px;border-bottom:1px solid #f1f5f9;vertical-align:middle}}
tr:last-child td{{border-bottom:none}}
tr:hover td{{background:#f8fafc}}
.badge{{display:inline-flex;align-items:center;gap:3px;padding:3px 8px;border-radius:99px;font-size:11px;font-weight:700;white-space:nowrap}}
.badge-success{{background:#dcfce7;color:#15803d}}
.badge-danger{{background:#fee2e2;color:#b91c1c}}
.badge-warning{{background:#fef3c7;color:#92400e}}
.badge-info{{background:#e0f2fe;color:#0369a1}}
.badge-gray{{background:#f1f5f9;color:#475569}}
.badge-alta{{background:#fee2e2;color:#b91c1c}}
.badge-media{{background:#fef3c7;color:#92400e}}
.badge-baixa{{background:#dbeafe;color:#1d4ed8}}
.badge-concluida{{background:#dcfce7;color:#15803d}}
.progress{{height:6px;background:#e2e8f0;border-radius:3px;overflow:hidden;width:90px;display:inline-block}}
.progress-bar{{height:100%;border-radius:3px;transition:width .3s}}
.filters{{display:flex;gap:10px;align-items:center;margin-bottom:16px;flex-wrap:wrap}}
.filters label{{font-size:12px;font-weight:600;color:var(--text-muted)}}
.filters select,.filters input[type=date],.filters input[type=text]{{padding:7px 10px;border:1.5px solid var(--border);border-radius:6px;font-size:13px;background:var(--card);color:var(--text);outline:none}}
.filters select:focus,.filters input:focus{{border-color:var(--primary)}}
.team-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(280px,1fr));gap:14px;margin-bottom:18px}}
.team-card{{background:var(--card);border-radius:10px;border:1px solid var(--border);padding:16px 18px}}
.team-card h3{{font-size:14px;font-weight:700;margin-bottom:2px}}
.team-card .leader{{font-size:11px;color:var(--text-muted);margin-bottom:12px}}
.team-stat{{display:flex;align-items:center;justify-content:space-between;margin-bottom:8px}}
.team-stat span{{font-size:12px;color:var(--text-muted)}}
.team-stat strong{{font-size:15px;font-weight:800}}
.editable-cell{{background:transparent;border:none;font-size:13px;color:var(--text);width:100%;cursor:pointer;font-family:inherit}}
.editable-cell:focus{{outline:2px solid var(--primary);border-radius:4px;padding:2px 4px;background:#f0f7ff}}
input[type=date].editable-cell{{cursor:pointer}}
.modal-overlay{{display:none;position:fixed;inset:0;background:rgba(0,0,0,.45);z-index:200;align-items:center;justify-content:center}}
.modal-overlay.open{{display:flex}}
.modal{{background:var(--card);border-radius:12px;padding:24px;width:500px;max-width:95vw;max-height:90vh;overflow-y:auto}}
.modal h2{{font-size:16px;font-weight:700;margin-bottom:16px}}
.form-group{{margin-bottom:13px}}
.form-group label{{display:block;font-size:11px;font-weight:700;color:var(--text-muted);margin-bottom:4px;text-transform:uppercase;letter-spacing:.4px}}
.form-group input,.form-group select,.form-group textarea{{width:100%;padding:8px 10px;border:1.5px solid var(--border);border-radius:6px;font-size:13px;color:var(--text);outline:none;background:var(--card);font-family:inherit}}
.form-group input:focus,.form-group select:focus,.form-group textarea:focus{{border-color:var(--primary)}}
.form-group textarea{{resize:vertical;min-height:70px}}
.modal-footer{{display:flex;gap:10px;justify-content:flex-end;margin-top:18px}}
.chart-wrap{{padding:14px;height:240px}}
.chart-wrap-lg{{padding:14px;height:300px}}
.empty{{text-align:center;padding:36px;color:var(--text-muted)}}
.empty .icon{{font-size:32px;margin-bottom:8px}}
.empty p{{font-size:13px}}
#toast{{position:fixed;bottom:24px;right:24px;background:#1e293b;color:#fff;padding:12px 18px;border-radius:8px;font-size:13px;font-weight:500;z-index:999;display:none;box-shadow:0 4px 12px rgba(0,0,0,.2)}}
::-webkit-scrollbar{{width:6px;height:6px}}
::-webkit-scrollbar-thumb{{background:#cbd5e1;border-radius:3px}}
.person-name{{font-weight:600}}
.person-sub{{font-size:11px;color:var(--text-muted)}}
.two-col{{display:grid;grid-template-columns:1fr 1fr;gap:16px}}
@media(max-width:900px){{.two-col{{grid-template-columns:1fr}}}}
.activity-list{{list-style:none;padding:0}}
.activity-list li{{padding:6px 14px;border-bottom:1px solid #f1f5f9;font-size:12.5px;color:var(--text)}}
.activity-list li:before{{content:"•";color:var(--primary);margin-right:6px}}
.activity-list li:last-child{{border-bottom:none}}
.overdue{{background:#fff5f5!important}}
</style>
</head>
<body>

<nav id="sidebar">
  <div class="logo">
    <span>Convergint</span>
    <small>Dashboard RDP</small>
  </div>
  <nav>
    <a href="#" class="active" onclick="go('overview',this)"><span class="icon">📊</span> Visão Geral</a>
    <a href="#" onclick="go('daily',this)"><span class="icon">📅</span> Status Diário</a>
    <a href="#" onclick="go('teams',this)"><span class="icon">👥</span> Por Equipe</a>
    <a href="#" onclick="go('projects',this)"><span class="icon">🏗️</span> Projetos</a>
    <a href="#" onclick="go('pending',this)"><span class="icon">⚠️</span> Pendências</a>
    <a href="#" onclick="go('history',this)"><span class="icon">📈</span> Histórico</a>
  </nav>
  <div class="sidebar-footer">
    Atualizado em<br><strong id="gen-date">{generated_at}</strong><br><br>
    Para recarregar os dados,<br>execute <strong>gerar_dashboard.bat</strong>
  </div>
</nav>

<div id="main">
  <div id="topbar">
    <h1 id="topbar-title">Visão Geral</h1>
    <div class="actions">
      <input type="date" id="global-date" style="padding:6px 10px;border:1.5px solid var(--border);border-radius:6px;font-size:13px;background:#fff">
      <button class="btn btn-outline" onclick="exportStatus()">📤 Exportar</button>
    </div>
  </div>

  <div id="content">

    <!-- VISÃO GERAL -->
    <div id="section-overview" class="section active">
      <div class="section-title">Visão Geral</div>
      <div class="section-sub" id="overview-sub"></div>
      <div class="kpi-grid">
        <div class="kpi-card" style="--accent:#0066CC"><div class="label">Colaboradores</div><div class="value" id="k-total">—</div><div class="sub">cadastrados</div></div>
        <div class="kpi-card" style="--accent:#16a34a"><div class="label">Enviaram RDP</div><div class="value" id="k-sent">—</div><div class="sub" id="k-sent-sub">hoje</div></div>
        <div class="kpi-card" style="--accent:#dc2626"><div class="label">Pendentes</div><div class="value" id="k-pend">—</div><div class="sub" id="k-pend-sub">não enviaram</div></div>
        <div class="kpi-card" style="--accent:#d97706"><div class="label">Compliance</div><div class="value" id="k-comp">—</div><div class="sub">taxa de envio</div></div>
        <div class="kpi-card" style="--accent:#0891b2"><div class="label">Projetos Ativos</div><div class="value" id="k-proj">—</div><div class="sub">no período</div></div>
        <div class="kpi-card" style="--accent:#7c3aed"><div class="label">Pendências</div><div class="value" id="k-pend-open">—</div><div class="sub">em aberto</div></div>
      </div>
      <div class="two-col">
        <div class="card"><div class="card-header"><h2>Compliance por Equipe</h2></div><div class="chart-wrap"><canvas id="chart-teams"></canvas></div></div>
        <div class="card"><div class="card-header"><h2>Envios — últimos 14 dias</h2></div><div class="chart-wrap"><canvas id="chart-daily"></canvas></div></div>
      </div>
      <div class="card">
        <div class="card-header"><h2>⚠️ Pendentes na data selecionada</h2><span class="badge badge-danger" id="ov-pend-cnt">0</span></div>
        <div class="card-body"><table><thead><tr><th>Colaborador</th><th>Equipe</th><th>Último Envio</th><th>Dias em Atraso</th></tr></thead><tbody id="ov-pend-table"></tbody></table></div>
      </div>
    </div>

    <!-- STATUS DIÁRIO -->
    <div id="section-daily" class="section">
      <div class="section-title">Status Diário</div>
      <div class="section-sub">Envio de RDP por colaborador na data selecionada.</div>
      <div class="filters">
        <label>Data:</label><input type="date" id="daily-date" oninput="renderDaily()">
        <label>Equipe:</label><select id="daily-team" onchange="renderDaily()"><option value="">Todas</option></select>
        <label>Status:</label>
        <select id="daily-status" onchange="renderDaily()">
          <option value="">Todos</option><option value="sent">Enviou</option><option value="pend">Pendente</option>
        </select>
        <input type="text" id="daily-search" placeholder="🔍 Buscar..." oninput="renderDaily()" style="min-width:180px">
      </div>
      <div class="card">
        <div class="card-header">
          <h2 id="daily-title">Colaboradores</h2>
          <div style="display:flex;gap:8px">
            <span class="badge badge-success" id="d-sent-cnt">0</span>
            <span class="badge badge-danger" id="d-pend-cnt">0</span>
          </div>
        </div>
        <div class="card-body">
          <table><thead><tr><th>Status</th><th>Colaborador</th><th>Equipe</th><th>Projeto(s)</th><th>Disciplina</th><th>Horário</th><th>Arquivo RDP</th></tr></thead>
          <tbody id="daily-tbody"></tbody></table>
        </div>
      </div>
    </div>

    <!-- POR EQUIPE -->
    <div id="section-teams" class="section">
      <div class="section-title">Performance por Equipe</div>
      <div class="section-sub">KPI de compliance individual e coletivo por período.</div>
      <div class="filters">
        <label>De:</label><input type="date" id="team-from">
        <span style="font-size:12px;color:var(--text-muted)">até</span>
        <input type="date" id="team-to">
        <button class="btn btn-primary btn-sm" onclick="renderTeams()">Filtrar</button>
      </div>
      <div class="team-grid" id="team-grid"></div>
      <div class="card">
        <div class="card-header"><h2>Ranking Individual</h2></div>
        <div class="card-body">
          <table><thead><tr><th>#</th><th>Colaborador</th><th>Equipe</th><th>Dias Enviados</th><th>Dias Esperados</th><th>Compliance</th><th>Último Envio</th></tr></thead>
          <tbody id="ranking-tbody"></tbody></table>
        </div>
      </div>
    </div>

    <!-- PROJETOS -->
    <div id="section-projects" class="section">
      <div class="section-title">Projetos</div>
      <div class="section-sub">Atividades e RDPs agrupados por projeto e disciplina.</div>
      <div class="filters">
        <input type="text" id="proj-search" placeholder="🔍 Buscar projeto..." oninput="renderProjects()" style="min-width:200px">
        <label>Disciplina:</label>
        <select id="proj-disc" onchange="renderProjects()">
          <option value="">Todas</option><option>AV</option><option>BMS</option><option>SDAI</option><option>SEG</option>
        </select>
      </div>
      <div id="projects-container"></div>
    </div>

    <!-- PENDÊNCIAS -->
    <div id="section-pending" class="section">
      <div class="section-title">Pendências</div>
      <div class="section-sub">Itens extraídos automaticamente dos RDPs e adicionados manualmente. As alterações são salvas no arquivo <code>.pendencias.json</code>.</div>
      <div style="display:flex;gap:10px;align-items:center;margin-bottom:14px;flex-wrap:wrap">
        <button class="btn btn-primary btn-sm" onclick="abrirModal()">+ Nova Pendência</button>
        <select id="pf-status" onchange="renderPending()"><option value="">Todos os status</option><option>Aberta</option><option>Concluída</option></select>
        <select id="pf-priority" onchange="renderPending()"><option value="">Todas as prioridades</option><option>Alta</option><option>Média</option><option>Baixa</option></select>
        <select id="pf-proj" onchange="renderPending()"><option value="">Todos os projetos</option></select>
        <input type="text" id="pf-search" placeholder="🔍 Buscar..." oninput="renderPending()" style="min-width:180px;padding:7px 10px;border:1.5px solid var(--border);border-radius:6px;font-size:13px">
        <button class="btn btn-outline btn-sm" onclick="exportPending()">📤 Exportar</button>
        <button class="btn btn-outline btn-sm" onclick="salvarPendencias()" style="color:var(--success);border-color:var(--success)">💾 Salvar Pendências</button>
      </div>
      <div class="card">
        <div class="card-header">
          <h2>Lista de Pendências</h2>
          <div style="display:flex;gap:6px">
            <span class="badge badge-alta" id="cnt-alta">0 Alta</span>
            <span class="badge badge-media" id="cnt-media">0 Média</span>
            <span class="badge badge-baixa" id="cnt-baixa">0 Baixa</span>
          </div>
        </div>
        <div class="card-body">
          <table><thead><tr>
            <th>Status</th><th>Prioridade</th><th>Descrição</th><th>Projeto</th>
            <th>Responsável</th><th>Criado em</th><th>Previsão Término</th>
            <th>Conclusão Real</th><th>Origem</th><th>Ações</th>
          </tr></thead><tbody id="pend-tbody"></tbody></table>
        </div>
      </div>
    </div>

    <!-- HISTÓRICO -->
    <div id="section-history" class="section">
      <div class="section-title">Histórico</div>
      <div class="section-sub">Tendência de compliance e log completo de envios.</div>
      <div class="filters">
        <label>De:</label><input type="date" id="hist-from">
        <label>Até:</label><input type="date" id="hist-to">
        <button class="btn btn-primary btn-sm" onclick="renderHistory()">Filtrar</button>
      </div>
      <div class="card"><div class="card-header"><h2>Compliance Diário</h2></div><div class="chart-wrap-lg"><canvas id="chart-hist"></canvas></div></div>
      <div class="card">
        <div class="card-header"><h2>Log de Envios</h2></div>
        <div class="card-body">
          <table><thead><tr><th>Data</th><th>Colaborador</th><th>Equipe</th><th>Projeto</th><th>Disciplina</th><th>Arquivo</th><th>Horário</th></tr></thead>
          <tbody id="hist-tbody"></tbody></table>
        </div>
      </div>
    </div>

  </div>
</div>

<!-- Modal Pendência -->
<div class="modal-overlay" id="modal">
  <div class="modal">
    <h2 id="modal-h">Nova Pendência</h2>
    <div class="form-group"><label>Descrição *</label><textarea id="m-desc" rows="3"></textarea></div>
    <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px">
      <div class="form-group"><label>Projeto</label><select id="m-proj"><option value="">—</option></select></div>
      <div class="form-group"><label>Disciplina</label><select id="m-disc"><option value="">—</option><option>AV</option><option>BMS</option><option>SDAI</option><option>SEG</option></select></div>
      <div class="form-group"><label>Prioridade *</label><select id="m-priority"><option>Alta</option><option selected>Média</option><option>Baixa</option></select></div>
      <div class="form-group"><label>Status</label><select id="m-status"><option>Aberta</option><option>Concluída</option></select></div>
      <div class="form-group"><label>Responsável</label><select id="m-resp"><option value="">—</option></select></div>
      <div class="form-group"><label>Previsão de Término</label><input type="date" id="m-forecast"></div>
      <div class="form-group"><label>Conclusão Real</label><input type="date" id="m-done"></div>
      <div class="form-group"><label>Origem</label><input type="text" id="m-origin" placeholder="Ex: RDP João — 22/04/2026"></div>
    </div>
    <div class="modal-footer">
      <button class="btn btn-outline" onclick="fecharModal()">Cancelar</button>
      <button class="btn btn-primary" onclick="salvarItem()">Salvar</button>
    </div>
  </div>
</div>

<div id="toast"></div>

<script>
// ── Dados embutidos ──────────────────────────────────────────────────────────
const RAW = {data_json};
const state = {{
  pessoas:    RAW.pessoas    || [],
  emails:     RAW.emails     || [],
  docs:       RAW.docs       || [],
  pendencias: RAW.pendencias || [],
  editingId:  null,
  charts:     {{}}
}};

// ── Init ─────────────────────────────────────────────────────────────────────
document.addEventListener('DOMContentLoaded', () => {{
  const today = RAW.today || new Date().toISOString().split('T')[0];
  document.getElementById('global-date').value  = today;
  document.getElementById('daily-date').value   = today;
  const d30 = new Date(); d30.setDate(d30.getDate()-30);
  const d30s = d30.toISOString().split('T')[0];
  document.getElementById('team-from').value = d30s;
  document.getElementById('team-to').value   = today;
  document.getElementById('hist-from').value = d30s;
  document.getElementById('hist-to').value   = today;

  document.getElementById('global-date').addEventListener('change', () => {{
    document.getElementById('daily-date').value = document.getElementById('global-date').value;
    renderOverview();
  }});

  populateFilters();
  renderOverview();
}});

// ── Helpers ───────────────────────────────────────────────────────────────────
const fmt = d => d ? d.split('-').reverse().join('/') : '—';
const fmtTime = iso => {{ const m=(iso||'').match(/T(\\d{{2}}:\\d{{2}})/); return m?m[1]:'—'; }};
const getDate = () => document.getElementById('global-date').value;
const teamShort = t => (t||'').replace('PRG BRA - ','').replace('BMS Desenvolvimento Telas','BMS Dev.Telas');

function getSentEmails(date) {{
  const s = new Set();
  state.emails.filter(e => e.dataStr === date).forEach(e => {{
    if (e.emailRemetente) s.add(e.emailRemetente);
    if (e.parsed) {{
      const prof = (e.parsed.professional||'').toLowerCase();
      state.pessoas.forEach(p => {{
        if (p.nome.toLowerCase().startsWith(prof.split(' ')[0])) s.add(p.email);
      }});
    }}
  }});
  // também considera docs escaneados (para quem não enviou e-mail mas tem doc)
  state.docs.filter(d => d.folderDate === date).forEach(d => {{
    if (d.profEmail) s.add(d.profEmail);
  }});
  return s;
}}

function getUniqueDates() {{
  const s = new Set([
    ...state.emails.map(e => e.dataStr),
    ...state.docs.map(d => d.folderDate)
  ].filter(Boolean));
  return [...s].sort();
}}

function getProjects() {{
  const s = new Set([
    ...state.emails.map(e => e.parsed?.project).filter(Boolean),
    ...state.docs.map(d => d.project).filter(Boolean)
  ]);
  return [...s].sort();
}}

function workDays(from, to) {{
  let c=0, d=new Date(from), e=new Date(to);
  while(d<=e){{ if(d.getDay()!==0&&d.getDay()!==6) c++; d.setDate(d.getDate()+1); }}
  return c||1;
}}

function destroyChart(k) {{ if(state.charts[k]){{ state.charts[k].destroy(); delete state.charts[k]; }} }}

// ── Navigation ────────────────────────────────────────────────────────────────
function go(id, el) {{
  document.querySelectorAll('.section').forEach(s=>s.classList.remove('active'));
  document.querySelectorAll('#sidebar nav a').forEach(a=>a.classList.remove('active'));
  document.getElementById('section-'+id).classList.add('active');
  if(el) el.classList.add('active');
  const t={{overview:'Visão Geral',daily:'Status Diário',teams:'Por Equipe',projects:'Projetos',pending:'Pendências',history:'Histórico'}};
  document.getElementById('topbar-title').textContent = t[id]||id;
  if(id==='teams') renderTeams();
  if(id==='projects') renderProjects();
  if(id==='pending') renderPending();
  if(id==='history') renderHistory();
  if(id==='daily') renderDaily();
  return false;
}}

// ── Populate filters ──────────────────────────────────────────────────────────
function populateFilters() {{
  const teams = [...new Set(state.pessoas.map(p=>p.equipe))].sort();
  document.getElementById('daily-team').innerHTML =
    '<option value="">Todas</option>' + teams.map(t=>`<option value="${{t}}">${{teamShort(t)}}</option>`).join('');

  const projs = getProjects();
  ['pf-proj','m-proj'].forEach(id => {{
    const el = document.getElementById(id);
    if(el) el.innerHTML = '<option value="">—</option>' + projs.map(p=>`<option>${{p}}</option>`).join('');
  }});
  const respEl = document.getElementById('m-resp');
  if(respEl) respEl.innerHTML = '<option value="">—</option>' + state.pessoas.map(p=>`<option value="${{p.email}}">${{p.nome}}</option>`).join('');
}}

// ── Overview ──────────────────────────────────────────────────────────────────
function renderOverview() {{
  const date  = getDate();
  const sent  = getSentEmails(date);
  const total = state.pessoas.length;
  const sentN = state.pessoas.filter(p=>sent.has(p.email)).length;
  const pend  = total - sentN;
  const comp  = total ? Math.round(sentN/total*100) : 0;
  const projs = getProjects().length;
  const pOpen = state.pendencias.filter(p=>p.status==='Aberta').length;

  document.getElementById('k-total').textContent   = total||'—';
  document.getElementById('k-sent').textContent    = total?sentN:'—';
  document.getElementById('k-pend').textContent    = total?pend:'—';
  document.getElementById('k-comp').textContent    = total?comp+'%':'—';
  document.getElementById('k-proj').textContent    = projs||'—';
  document.getElementById('k-pend-open').textContent = pOpen;
  document.getElementById('k-sent-sub').textContent  = `em ${{fmt(date)}}`;
  document.getElementById('k-pend-sub').textContent  = `em ${{fmt(date)}}`;
  document.getElementById('overview-sub').textContent = `Referência: ${{fmt(date)}} · Dados atualizados em: ${{RAW.generatedAt||'—'}}`;

  renderChartTeams(date);
  renderChartDailyLine();
  renderPendingToday(date, sent);
}}

function renderChartTeams(date) {{
  destroyChart('teams');
  const teams = [...new Set(state.pessoas.map(p=>p.equipe))];
  const sent  = getSentEmails(date);
  const pal   = ['#0066CC','#16a34a','#d97706','#dc2626','#7c3aed','#0891b2','#ea580c'];
  const labels=[],data=[],colors=[];
  teams.forEach((t,i) => {{
    const members = state.pessoas.filter(p=>p.equipe===t);
    const r = members.length ? Math.round(members.filter(m=>sent.has(m.email)).length/members.length*100) : 0;
    labels.push(teamShort(t).substring(0,22)); data.push(r); colors.push(pal[i%pal.length]);
  }});
  const ctx = document.getElementById('chart-teams').getContext('2d');
  state.charts.teams = new Chart(ctx,{{
    type:'bar', data:{{labels,datasets:[{{label:'%',data,backgroundColor:colors,borderRadius:5}}]}},
    options:{{responsive:true,maintainAspectRatio:false,plugins:{{legend:{{display:false}}}},scales:{{y:{{min:0,max:100,ticks:{{callback:v=>v+'%'}}}}}}}}
  }});
}}

function renderChartDailyLine() {{
  destroyChart('daily');
  const dates = getUniqueDates().slice(-14);
  if(!dates.length) return;
  const total = state.pessoas.length||1;
  const data  = dates.map(d => Math.round(state.pessoas.filter(p=>getSentEmails(d).has(p.email)).length/total*100));
  const ctx   = document.getElementById('chart-daily').getContext('2d');
  state.charts.daily = new Chart(ctx,{{
    type:'line',
    data:{{labels:dates.map(fmt),datasets:[{{label:'%',data,borderColor:'#0066CC',backgroundColor:'rgba(0,102,204,.08)',fill:true,tension:.4,pointRadius:4,pointBackgroundColor:'#0066CC'}}]}},
    options:{{responsive:true,maintainAspectRatio:false,plugins:{{legend:{{display:false}}}},scales:{{y:{{min:0,max:100,ticks:{{callback:v=>v+'%'}}}}}}}}
  }});
}}

function renderPendingToday(date, sent) {{
  const pendPeople = state.pessoas.filter(p=>!sent.has(p.email));
  document.getElementById('ov-pend-cnt').textContent = pendPeople.length;
  const tbody = document.getElementById('ov-pend-table');
  if(!pendPeople.length) {{
    tbody.innerHTML=`<tr><td colspan="4" class="empty"><div class="icon">✅</div><p>${{state.pessoas.length?'Todos enviaram!':'Carregue os dados.'}}</p></td></tr>`;
    return;
  }}
  tbody.innerHTML = pendPeople.map(p => {{
    const last = [...state.emails].reverse().find(e=>e.emailRemetente===p.email);
    const lastD = last?.dataStr||null;
    const days  = lastD ? Math.round((new Date(date)-new Date(lastD))/86400000) : '∞';
    const cls   = days==='∞'||days>3?'badge-alta':days>1?'badge-warning':'badge-info';
    return `<tr>
      <td><div class="person-name">${{p.nome}}</div></td>
      <td><span class="badge badge-gray">${{teamShort(p.equipe)}}</span></td>
      <td>${{lastD?fmt(lastD):'<span style="color:var(--text-muted)">Nunca enviou</span>'}}</td>
      <td><span class="badge ${{cls}}">${{days==='∞'?'Sem histórico':days+' dia(s)'}}</span></td>
    </tr>`;
  }}).join('');
}}

// ── Status Diário ─────────────────────────────────────────────────────────────
function renderDaily() {{
  const date   = document.getElementById('daily-date').value||getDate();
  const team   = document.getElementById('daily-team').value;
  const status = document.getElementById('daily-status').value;
  const search = document.getElementById('daily-search').value.toLowerCase();
  const sent   = getSentEmails(date);

  let rows = state.pessoas.filter(p =>
    (!team   || p.equipe===team) &&
    (!search || p.nome.toLowerCase().includes(search))
  ).map(p => {{
    const isSent  = sent.has(p.email);
    const myEmails = state.emails.filter(e => e.dataStr===date && (e.emailRemetente===p.email || (e.parsed && p.nome.toLowerCase().startsWith((e.parsed.professional||'').toLowerCase().split(' ')[0]))));
    const myDocs   = state.docs.filter(d => d.folderDate===date && d.profEmail===p.email);
    const projects = [...new Set([...myEmails.map(e=>e.parsed?.project),...myDocs.map(d=>d.project)].filter(Boolean))].join(', ')||'—';
    const discs    = [...new Set([...myEmails.map(e=>e.parsed?.discipline),...myDocs.map(d=>d.discipline)].filter(Boolean))].join(', ')||'—';
    const arquivo  = myEmails[0]?.arquivo || myDocs[0]?.filename || '—';
    const timeRaw  = myEmails[0]?.dataRecebimento||'';
    return {{p, isSent, projects, discs, arquivo, timeRaw}};
  }});

  if(status==='sent')  rows = rows.filter(r=>r.isSent);
  if(status==='pend')  rows = rows.filter(r=>!r.isSent);

  const sentN = rows.filter(r=>r.isSent).length;
  document.getElementById('daily-title').textContent = `Colaboradores — ${{fmt(date)}}`;
  document.getElementById('d-sent-cnt').textContent  = `${{sentN}} enviaram`;
  document.getElementById('d-pend-cnt').textContent  = `${{rows.length-sentN}} pendentes`;

  document.getElementById('daily-tbody').innerHTML = rows.length
    ? rows.map(r=>`<tr>
        <td><span class="badge ${{r.isSent?'badge-success':'badge-danger'}}">${{r.isSent?'✓ Enviou':'✗ Pendente'}}</span></td>
        <td><div class="person-name">${{r.p.nome}}</div><div class="person-sub">${{r.p.email}}</div></td>
        <td><span class="badge badge-gray">${{teamShort(r.p.equipe)}}</span></td>
        <td>${{r.projects!=='—'?`<span class="badge badge-info">${{r.projects}}</span>`:'—'}}</td>
        <td>${{r.discs!=='—'?`<span class="badge badge-gray">${{r.discs}}</span>`:'—'}}</td>
        <td style="font-variant-numeric:tabular-nums">${{r.timeRaw?fmtTime(r.timeRaw):'—'}}</td>
        <td style="max-width:220px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;font-size:11px;color:var(--text-muted)" title="${{r.arquivo}}">${{r.arquivo!=='—'?'📎 '+r.arquivo:'—'}}</td>
      </tr>`).join('')
    : `<tr><td colspan="7" class="empty"><div class="icon">📋</div><p>Nenhum colaborador encontrado.</p></td></tr>`;
}}

// ── Equipes ───────────────────────────────────────────────────────────────────
function renderTeams() {{
  const from  = document.getElementById('team-from').value;
  const to    = document.getElementById('team-to').value;
  const dates = getUniqueDates().filter(d=>(!from||d>=from)&&(!to||d<=to));
  const expD  = workDays(from||dates[0]||'2026-01-01', to||dates[dates.length-1]||'2026-12-31');
  const teams = [...new Set(state.pessoas.map(p=>p.equipe))].sort();
  const pal   = ['#0066CC','#16a34a','#d97706','#dc2626','#7c3aed','#0891b2'];

  document.getElementById('team-grid').innerHTML = teams.map((t,i) => {{
    const members = state.pessoas.filter(p=>p.equipe===t);
    let totalSent=0;
    members.forEach(m => {{
      totalSent += dates.filter(d => getSentEmails(d).has(m.email)).length;
    }});
    const exp  = members.length * expD;
    const rate = exp>0 ? Math.round(totalSent/exp*100) : 0;
    const col  = pal[i%pal.length];
    const cls  = rate>=80?'var(--success)':rate>=50?'var(--warning)':'var(--danger)';
    const name = teamShort(t);
    return `<div class="team-card">
      <h3>${{name.split(' - ')[0]}}</h3>
      <div class="leader">Líder: ${{name.split(' - ').pop()}}</div>
      <div class="team-stat"><span>Membros</span><strong>${{members.length}}</strong></div>
      <div class="team-stat"><span>Compliance (período)</span><strong style="color:${{cls}}">${{rate}}%</strong></div>
      <div style="margin-top:8px">
        <div style="display:flex;justify-content:space-between;font-size:11px;color:var(--text-muted);margin-bottom:4px">
          <span>${{totalSent}} envios</span><span>de ${{exp}} esperados</span>
        </div>
        <div class="progress" style="width:100%"><div class="progress-bar" style="width:${{rate}}%;background:${{col}}"></div></div>
      </div>
    </div>`;
  }}).join('');

  const ranking = state.pessoas.map(p => {{
    const sentD = dates.filter(d=>getSentEmails(d).has(p.email)).length;
    const comp  = expD>0 ? Math.round(sentD/expD*100) : 0;
    const last  = [...state.emails].reverse().find(e=>e.emailRemetente===p.email);
    return {{p, sentD, expD, comp, lastDate:last?.dataStr}};
  }}).sort((a,b)=>b.comp-a.comp);

  document.getElementById('ranking-tbody').innerHTML = ranking.map((r,i) => {{
    const cls = r.comp>=80?'var(--success)':r.comp>=50?'var(--warning)':'var(--danger)';
    return `<tr>
      <td style="font-weight:700;color:var(--text-muted)">${{i+1}}</td>
      <td><div class="person-name">${{r.p.nome}}</div></td>
      <td><span class="badge badge-gray">${{teamShort(r.p.equipe)}}</span></td>
      <td>${{r.sentD}}</td><td>${{r.expD}}</td>
      <td><div style="display:flex;align-items:center;gap:8px">
        <div class="progress"><div class="progress-bar" style="width:${{r.comp}}%;background:${{cls}}"></div></div>
        <span style="font-size:12px;font-weight:700;color:${{cls}}">${{r.comp}}%</span>
      </div></td>
      <td>${{r.lastDate?fmt(r.lastDate):'<span style="color:var(--text-muted)">—</span>'}}</td>
    </tr>`;
  }}).join('');
}}

// ── Projetos ──────────────────────────────────────────────────────────────────
function renderProjects() {{
  const search = document.getElementById('proj-search').value.toLowerCase();
  const disc   = document.getElementById('proj-disc').value;

  const projMap = {{}};
  // De e-mails
  state.emails.forEach(e => {{
    if(!e.parsed) return;
    const k = e.parsed.project;
    if(!projMap[k]) projMap[k] = {{rdps:[]}};
    projMap[k].rdps.push({{date:e.parsed.date,discipline:e.parsed.discipline,professional:e.parsed.professional,arquivo:e.arquivo,bodyPreview:e.bodyPreview,activities:[],pendingItems:[],hasText:false}});
  }});
  // De docs (sobrescreve com conteúdo extraído)
  state.docs.forEach(d => {{
    const k = d.project;
    if(!projMap[k]) projMap[k] = {{rdps:[]}};
    // verifica se já existe entrada do e-mail
    const existing = projMap[k].rdps.find(r=>r.arquivo===d.filename);
    if(existing) {{
      existing.activities   = d.activities;
      existing.pendingItems = d.pendingItems;
      existing.hasText      = d.hasText;
    }} else {{
      projMap[k].rdps.push({{date:d.date,discipline:d.discipline,professional:d.professional,arquivo:d.filename,bodyPreview:d.textPreview,activities:d.activities,pendingItems:d.pendingItems,hasText:d.hasText}});
    }}
  }});

  let entries = Object.entries(projMap);
  if(search) entries = entries.filter(([k])=>k.toLowerCase().includes(search));
  if(disc)   entries = entries.filter(([,v])=>v.rdps.some(r=>r.discipline===disc));

  const cont = document.getElementById('projects-container');
  if(!entries.length) {{
    cont.innerHTML=`<div class="empty"><div class="icon">🏗️</div><p>Nenhum projeto encontrado.</p></div>`;
    return;
  }}

  cont.innerHTML = entries.map(([proj,data]) => {{
    const rdps   = data.rdps;
    const discs  = [...new Set(rdps.map(r=>r.discipline))];
    const dates  = [...new Set(rdps.map(r=>r.date))].sort().reverse();
    const discB  = discs.map(d=>`<span class="badge ${{d==='AV'?'badge-info':d==='BMS'?'badge-warning':d==='SDAI'?'badge-gray':'badge-danger'}}">${{d}}</span>`).join(' ');
    const hasActs = rdps.some(r=>r.activities?.length);

    return `<div class="card">
      <div class="card-header" style="cursor:pointer" onclick="toggleProj('${{proj}}')">
        <div>
          <h2 style="font-size:16px">${{proj}} <span style="font-size:12px;color:var(--text-muted);font-weight:400">▼</span></h2>
          <div style="margin-top:4px;display:flex;gap:6px">${{discB}}</div>
        </div>
        <div style="text-align:right;font-size:12px;color:var(--text-muted)">
          <div>${{rdps.length}} RDPs</div>
          <div>Último: ${{fmt(dates[0]||'')}}</div>
        </div>
      </div>
      <div id="p-${{proj}}" style="display:none">
        <div class="card-body">
          <table><thead><tr><th>Data</th><th>Profissional</th><th>Disciplina</th><th>Atividades Extraídas</th><th>Pendências no Doc</th><th>Arquivo</th></tr></thead>
          <tbody>
            ${{rdps.sort((a,b)=>b.date.localeCompare(a.date)).map(r=>`<tr>
              <td>${{fmt(r.date)}}</td>
              <td>${{r.professional}}</td>
              <td><span class="badge badge-gray">${{r.discipline}}</span></td>
              <td style="max-width:250px">
                ${{r.activities?.length
                  ? `<ul class="activity-list">${{r.activities.slice(0,5).map(a=>`<li>${{a}}</li>`).join('')}}</ul>`
                  : '<span style="color:var(--text-muted);font-size:11px">—</span>'}}
              </td>
              <td style="max-width:200px">
                ${{r.pendingItems?.length
                  ? `<ul class="activity-list">${{r.pendingItems.slice(0,3).map(p=>`<li style="color:var(--danger)">${{p}}</li>`).join('')}}</ul>`
                  : '<span style="color:var(--text-muted);font-size:11px">—</span>'}}
              </td>
              <td style="font-size:11px;color:var(--text-muted);max-width:180px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap" title="${{r.arquivo}}">📎 ${{r.arquivo||'—'}}</td>
            </tr>`).join('')}}
          </tbody></table>
        </div>
      </div>
    </div>`;
  }}).join('');
}}

function toggleProj(id) {{
  const el = document.getElementById('p-'+id);
  if(el) el.style.display = el.style.display==='none'?'block':'none';
}}

// ── Pendências ────────────────────────────────────────────────────────────────
function renderPending() {{
  const sf = document.getElementById('pf-status').value;
  const pf = document.getElementById('pf-priority').value;
  const jf = document.getElementById('pf-proj').value;
  const se = document.getElementById('pf-search').value.toLowerCase();

  let items = [...state.pendencias];
  if(sf) items = items.filter(i=>i.status===sf);
  if(pf) items = items.filter(i=>i.priority===pf);
  if(jf) items = items.filter(i=>i.project===jf);
  if(se) items = items.filter(i=>(i.desc||'').toLowerCase().includes(se)||(i.responsible||'').toLowerCase().includes(se));

  document.getElementById('cnt-alta').textContent  = state.pendencias.filter(i=>i.status==='Aberta'&&i.priority==='Alta').length+' Alta';
  document.getElementById('cnt-media').textContent = state.pendencias.filter(i=>i.status==='Aberta'&&i.priority==='Média').length+' Média';
  document.getElementById('cnt-baixa').textContent = state.pendencias.filter(i=>i.status==='Aberta'&&i.priority==='Baixa').length+' Baixa';

  const today = new Date().toISOString().split('T')[0];
  const tbody = document.getElementById('pend-tbody');

  if(!items.length) {{
    tbody.innerHTML=`<tr><td colspan="10" class="empty"><div class="icon">⚠️</div><p>Nenhuma pendência. Clique em "+ Nova Pendência" para adicionar.</p></td></tr>`;
    return;
  }}

  tbody.innerHTML = items.map(item => {{
    const overdue = item.status==='Aberta' && item.forecast && item.forecast < today;
    const resp    = state.pessoas.find(p=>p.email===item.responsible)?.nome || item.responsible || '—';
    return `<tr class="${{overdue?'overdue':''}}">
      <td><select class="editable-cell" onchange="updPend(${{item.id}},'status',this.value)" style="width:100px">
        <option ${{item.status==='Aberta'?'selected':''}}>Aberta</option>
        <option ${{item.status==='Concluída'?'selected':''}}>Concluída</option>
      </select></td>
      <td><select class="editable-cell" onchange="updPend(${{item.id}},'priority',this.value)" style="width:75px">
        <option ${{item.priority==='Alta'?'selected':''}}>Alta</option>
        <option ${{item.priority==='Média'?'selected':''}}>Média</option>
        <option ${{item.priority==='Baixa'?'selected':''}}>Baixa</option>
      </select></td>
      <td style="min-width:200px">
        <div style="font-weight:600">${{item.desc}}</div>
        ${{overdue?'<span class="badge badge-alta" style="margin-top:3px">Atrasada</span>':''}}
        ${{item.autoDetected?'<span class="badge badge-gray" style="margin-top:3px;font-size:10px">Auto</span>':''}}
      </td>
      <td>${{item.project?`<span class="badge badge-info">${{item.project}}</span>`:'—'}}</td>
      <td style="font-size:12px">${{resp}}</td>
      <td style="font-size:12px">${{fmt(item.createdAt)}}</td>
      <td><input type="date" class="editable-cell" value="${{item.forecast||''}}" onchange="updPend(${{item.id}},'forecast',this.value)" style="width:128px;border:1px solid var(--border);border-radius:4px;padding:3px 6px;font-size:12px"></td>
      <td><input type="date" class="editable-cell" value="${{item.doneDate||''}}" onchange="updPend(${{item.id}},'doneDate',this.value)" style="width:128px;border:1px solid var(--border);border-radius:4px;padding:3px 6px;font-size:12px"></td>
      <td style="font-size:11px;color:var(--text-muted);max-width:150px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap" title="${{item.origin||''}}">${{item.origin||'—'}}</td>
      <td style="white-space:nowrap">
        <button class="btn btn-outline btn-sm" onclick="editItem(${{item.id}})" title="Editar">✏️</button>
        <button class="btn btn-sm btn-danger" onclick="delItem(${{item.id}})" title="Excluir">🗑️</button>
      </td>
    </tr>`;
  }}).join('');
}}

function updPend(id, field, value) {{
  const item = state.pendencias.find(p=>p.id===id);
  if(item) {{
    item[field] = value;
    if(field==='doneDate' && value) item.status='Concluída';
    renderPending();
    toast('Salvo ✓ — clique em "💾 Salvar Pendências" para persistir.');
  }}
}}

function abrirModal() {{
  state.editingId = null;
  document.getElementById('modal-h').textContent='Nova Pendência';
  ['m-desc','m-forecast','m-done','m-origin'].forEach(id=>{{ const el=document.getElementById(id); if(el) el.value=''; }});
  document.getElementById('m-proj').value='';
  document.getElementById('m-disc').value='';
  document.getElementById('m-priority').value='Média';
  document.getElementById('m-status').value='Aberta';
  document.getElementById('m-resp').value='';
  document.getElementById('modal').classList.add('open');
}}

function editItem(id) {{
  const item = state.pendencias.find(p=>p.id===id);
  if(!item) return;
  state.editingId = id;
  document.getElementById('modal-h').textContent='Editar Pendência';
  document.getElementById('m-desc').value    = item.desc||'';
  document.getElementById('m-proj').value    = item.project||'';
  document.getElementById('m-disc').value    = item.discipline||'';
  document.getElementById('m-priority').value= item.priority||'Média';
  document.getElementById('m-status').value  = item.status||'Aberta';
  document.getElementById('m-resp').value    = item.responsible||'';
  document.getElementById('m-forecast').value= item.forecast||'';
  document.getElementById('m-done').value    = item.doneDate||'';
  document.getElementById('m-origin').value  = item.origin||'';
  document.getElementById('modal').classList.add('open');
}}

function salvarItem() {{
  const desc = document.getElementById('m-desc').value.trim();
  if(!desc) {{ alert('Informe a descrição.'); return; }}
  const today = new Date().toISOString().split('T')[0];
  const data  = {{
    desc, project:document.getElementById('m-proj').value,
    discipline:document.getElementById('m-disc').value,
    priority:document.getElementById('m-priority').value,
    status:document.getElementById('m-status').value,
    responsible:document.getElementById('m-resp').value,
    forecast:document.getElementById('m-forecast').value,
    doneDate:document.getElementById('m-done').value,
    origin:document.getElementById('m-origin').value
  }};
  if(state.editingId !== null) {{
    const item = state.pendencias.find(p=>p.id===state.editingId);
    if(item) Object.assign(item, data);
  }} else {{
    data.id = Date.now(); data.createdAt = today; data.autoDetected = false;
    state.pendencias.push(data);
  }}
  fecharModal();
  renderPending();
  toast('Pendência salva. Clique em "💾 Salvar Pendências" para persistir.');
}}

function delItem(id) {{
  if(!confirm('Excluir esta pendência?')) return;
  state.pendencias = state.pendencias.filter(p=>p.id!==id);
  renderPending();
  toast('Excluída. Clique em "💾 Salvar Pendências" para persistir.');
}}

function fecharModal() {{ document.getElementById('modal').classList.remove('open'); }}
document.getElementById('modal').addEventListener('click', e=>{{ if(e.target===document.getElementById('modal')) fecharModal(); }});

// Salvar pendências — escreve JSON via download para reimportar no script
function salvarPendencias() {{
  const json = JSON.stringify(state.pendencias, null, 2);
  const blob = new Blob([json], {{type:'application/json'}});
  const a    = document.createElement('a'); a.href = URL.createObjectURL(blob);
  a.download = '.pendencias.json'; a.click();
  toast('Baixe e salve o arquivo ".pendencias.json" na pasta do dashboard para persistir as pendências.');
}}

// ── Histórico ─────────────────────────────────────────────────────────────────
function renderHistory() {{
  const from  = document.getElementById('hist-from').value;
  const to    = document.getElementById('hist-to').value;
  const dates = getUniqueDates().filter(d=>(!from||d>=from)&&(!to||d<=to));
  const total = state.pessoas.length||1;

  destroyChart('hist');
  if(dates.length) {{
    const ctx = document.getElementById('chart-hist').getContext('2d');
    state.charts.hist = new Chart(ctx,{{
      type:'line',
      data:{{
        labels:dates.map(fmt),
        datasets:[{{label:'%',data:dates.map(d=>Math.round(state.pessoas.filter(p=>getSentEmails(d).has(p.email)).length/total*100)),borderColor:'#0066CC',backgroundColor:'rgba(0,102,204,.08)',fill:true,tension:.4,pointRadius:4,pointBackgroundColor:'#0066CC'}}]
      }},
      options:{{responsive:true,maintainAspectRatio:false,plugins:{{legend:{{display:false}}}},scales:{{y:{{min:0,max:100,ticks:{{callback:v=>v+'%'}}}}}}}}
    }});
  }}

  const allEmails = state.emails.filter(e=>(!from||e.dataStr>=from)&&(!to||e.dataStr<=to));
  const sorted = [...allEmails].sort((a,b)=>b.dataStr.localeCompare(a.dataStr));
  const tbody  = document.getElementById('hist-tbody');
  tbody.innerHTML = sorted.length
    ? sorted.map(e => {{
        const p = state.pessoas.find(x=>x.email===e.emailRemetente);
        return `<tr>
          <td>${{fmt(e.dataStr)}}</td>
          <td>${{e.nomeRemetente||p?.nome||e.emailRemetente}}</td>
          <td>${{p?`<span class="badge badge-gray">${{teamShort(p.equipe)}}</span>`:'—'}}</td>
          <td>${{e.parsed?.project?`<span class="badge badge-info">${{e.parsed.project}}</span>`:'—'}}</td>
          <td>${{e.parsed?.discipline?`<span class="badge badge-gray">${{e.parsed.discipline}}</span>`:'—'}}</td>
          <td style="font-size:11px;color:var(--text-muted);max-width:200px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap" title="${{e.arquivo}}">📎 ${{e.arquivo||'—'}}</td>
          <td style="font-variant-numeric:tabular-nums">${{fmtTime(e.dataRecebimento)}}</td>
        </tr>`;
      }}).join('')
    : `<tr><td colspan="7" class="empty"><p>Nenhum envio no período.</p></td></tr>`;
}}

// ── Export ────────────────────────────────────────────────────────────────────
function exportStatus() {{
  const date = getDate();
  const sent = getSentEmails(date);
  const rows = state.pessoas.map(p=>({{
    Nome:p.nome, Email:p.email, Equipe:p.equipe,
    Status:sent.has(p.email)?'Enviou':'Pendente', Data:date
  }}));
  const wb = XLSX.utils.book_new();
  XLSX.utils.book_append_sheet(wb, XLSX.utils.json_to_sheet(rows), 'Status RDP');
  if(state.pendencias.length) {{
    const pr = state.pendencias.map(p=>(({{Descrição:p.desc,Projeto:p.project,Prioridade:p.priority,Status:p.status,Responsável:state.pessoas.find(x=>x.email===p.responsible)?.nome||p.responsible,'Previsão':p.forecast,'Conclusão':p.doneDate,'Criado':p.createdAt}})));
    XLSX.utils.book_append_sheet(wb, XLSX.utils.json_to_sheet(pr), 'Pendências');
  }}
  XLSX.writeFile(wb, `RDP_${{date}}.xlsx`);
}}

function exportPending() {{
  if(!state.pendencias.length) {{ toast('Nenhuma pendência.'); return; }}
  const rows = state.pendencias.map(p=>(({{Descrição:p.desc,Projeto:p.project,Disciplina:p.discipline,Prioridade:p.priority,Status:p.status,Responsável:state.pessoas.find(x=>x.email===p.responsible)?.nome||p.responsible,'Previsão':p.forecast,'Conclusão':p.doneDate,'Criado':p.createdAt,Origem:p.origin}})));
  const wb = XLSX.utils.book_new();
  XLSX.utils.book_append_sheet(wb, XLSX.utils.json_to_sheet(rows), 'Pendências');
  XLSX.writeFile(wb, `Pendencias_${{new Date().toISOString().split('T')[0]}}.xlsx`);
}}

function toast(msg) {{
  const el = document.getElementById('toast');
  el.textContent=msg; el.style.display='block';
  setTimeout(()=>el.style.display='none', 4000);
}}
</script>
</body>
</html>"""

    OUTPUT_HTML.write_text(html, encoding="utf-8")
    print(f"  ✓ Dashboard principal gerado: {OUTPUT_HTML}")

def gerar_html_pendencias(pessoas, emails, docs, pendencias):
    """Gera o dashboard focado em pendências e resumo RDP."""
    today_str      = date.today().isoformat()
    generated_at   = datetime.now().strftime("%d/%m/%Y às %H:%M")

    # Monta objetos de projeto a partir dos docs + emails
    proj_map = {}

    for doc in docs:
        k = doc.get("project", "DESCONHECIDO")
        if k not in proj_map:
            proj_map[k] = {
                "codigo": k,
                "disciplinas": [],
                "ultimoRDP": doc.get("folderDate", ""),
                "profissional": doc.get("professional", ""),
                "atividades": "",
                "riscos": "",
                "arquivos": []
            }
        entry = proj_map[k]
        if doc.get("discipline") and doc["discipline"] not in entry["disciplinas"]:
            entry["disciplinas"].append(doc["discipline"])
        if doc.get("folderDate", "") > entry.get("ultimoRDP", ""):
            entry["ultimoRDP"]    = doc["folderDate"]
            entry["profissional"] = doc.get("professional", entry["profissional"])
        if doc.get("filename"):
            entry["arquivos"].append(doc["filename"])
        # Concatena atividades e riscos extraídos
        if doc.get("activities"):
            entry["atividades"] += (" " if entry["atividades"] else "") + " ".join(doc["activities"][:5])
        if doc.get("pendingItems"):
            entry["riscos"] += (" " if entry["riscos"] else "") + " ".join(doc["pendingItems"][:3])

    # Complementa com dados de e-mails se projeto não tem doc
    for e in emails:
        if not e.get("parsed"): continue
        k = e["parsed"]["project"]
        if k not in proj_map:
            proj_map[k] = {
                "codigo": k,
                "disciplinas": [e["parsed"].get("discipline","")],
                "ultimoRDP": e.get("dataStr",""),
                "profissional": e["parsed"].get("professional",""),
                "atividades": e.get("bodyPreview","")[:300] if e.get("bodyPreview") else "",
                "riscos": "",
                "arquivos": [e.get("arquivo","")]
            }

    projetos_json = list(proj_map.values())

    # Converte pendências para formato do dashboard_pendencias.html
    pend_json = []
    for p in pendencias:
        pend_json.append({
            "id":          p.get("id", 0),
            "projeto":     p.get("project", ""),
            "disciplina":  p.get("discipline", ""),
            "prioridade":  p.get("priority", "Média"),
            "status":      p.get("status", "Aberta"),
            "desc":        p.get("desc", ""),
            "responsavel": p.get("responsible", ""),
            "forecast":    p.get("forecast", ""),
            "doneDate":    p.get("doneDate", ""),
            "origem":      p.get("origin", ""),
            "obs":         p.get("obs", ""),
            "createdAt":   p.get("createdAt", today_str),
            "autoDetected": p.get("autoDetected", False)
        })

    data_json = json.dumps({
        "projetos":    projetos_json,
        "pendencias":  pend_json,
        "generatedAt": generated_at,
        "today":       today_str
    }, ensure_ascii=False, default=str)

    # Lê o template HTML e injeta os dados
    template = OUTPUT_PENDENCIAS.read_text(encoding="utf-8")
    injected = template.replace(
        "const EMBEDDED = window.__RDP_DATA__ || null;",
        f"const EMBEDDED = {data_json};"
    )
    OUTPUT_PENDENCIAS.write_text(injected, encoding="utf-8")
    print(f"  ✓ Dashboard pendências gerado: {OUTPUT_PENDENCIAS}")

# ── Main ─────────────────────────────────────────────────────────────────────
def main():
    print("=" * 55)
    print("  Dashboard RDP — Convergint")
    print(f"  {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    print("=" * 55)

    print("\n[1/5] Carregando colaboradores...")
    pessoas = load_pessoas()

    print("\n[2/5] Carregando e-mails RDP...")
    emails = load_emails()

    print("\n[3/5] Escaneando documentos nas pastas RDP...")
    docs = scan_rdp_folders(pessoas)

    print("\n[4/5] Atualizando pendências...")
    pendencias = load_pendencias()
    pendencias = auto_update_pendencias(pendencias, docs)
    save_pendencias(pendencias)
    print(f"  ✓ {len(pendencias)} pendências ({sum(1 for p in pendencias if p['status']=='Aberta')} abertas)")

    print("\n[5/5] Gerando dashboards HTML...")
    gerar_html(pessoas, emails, docs, pendencias)
    if OUTPUT_PENDENCIAS.exists():
        gerar_html_pendencias(pessoas, emails, docs, pendencias)

    print("\n" + "=" * 55)
    print("  ✅ CONCLUÍDO!")
    print(f"  Dashboard principal:   {OUTPUT_HTML}")
    print(f"  Dashboard pendências:  {OUTPUT_PENDENCIAS}")
    print("=" * 55)

    # Abre ambos no navegador padrão
    import webbrowser
    webbrowser.open(str(OUTPUT_PENDENCIAS))
    webbrowser.open(str(OUTPUT_HTML))

if __name__ == "__main__":
    main()
